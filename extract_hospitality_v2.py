from docx import Document
import json

# Load the hospitality DOCX file
doc = Document('public/extradata/Combined file - Hospitality and Tourism.docx')

# Extract all text from tables
hospitality_data = {}

for table_idx, table in enumerate(doc.tables):
    print(f"\n=== TABLE {table_idx} ===")
    rows = table.rows
    
    for row_idx, row in enumerate(rows):
        cells = row.cells
        row_text = [cell.text.strip() for cell in cells]
        print(f"Row {row_idx}: {row_text}")
        
        if row_idx > 20:  # Limit output
            break

print("\n\nTotal tables:", len(doc.tables))
