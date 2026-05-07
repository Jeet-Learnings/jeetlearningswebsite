#!/usr/bin/env python3
"""
Script to extract "Top Institutions" from DOCX files and replace with "Where to Study?" sections.
Handles multiple careers per document (20+).
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Directory containing the DOCX files
DOCX_DIR = "public/extradata"

def extract_institutions_section(text):
    """
    Extract the "Top Institutions" section from text.
    Returns the extracted institutions text or None if not found.
    """
    # Pattern to match "Top Institutions" section
    pattern = r'Top Institutions\s*(.*?)(?=\n(?:Career Opportunities|Challenges|Future|Skills to Build|$))'
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    
    if match:
        return match.group(1).strip()
    return None

def replace_top_institutions_with_where_to_study(docx_path, output_path):
    """
    Read a DOCX file, find all "Top Institutions" sections, and replace with "Where to Study?" sections.
    """
    try:
        doc = Document(docx_path)
        replacements_made = 0
        
        # Extract all text to find "Top Institutions" sections
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        # Find all "Top Institutions" sections
        pattern = r'Top Institutions'
        
        # Process paragraphs
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            
            if 'Top Institutions' in para.text:
                # Replace "Top Institutions" with "Where to Study?"
                para.text = para.text.replace('Top Institutions', 'Where to Study?')
                replacements_made += 1
                print(f"  ✓ Replaced 'Top Institutions' with 'Where to Study?' at paragraph {i}")
            
            i += 1
        
        # Save the modified document
        doc.save(output_path)
        print(f"✓ Processed: {os.path.basename(docx_path)}")
        print(f"  Replacements made: {replacements_made}")
        return replacements_made
        
    except Exception as e:
        print(f"✗ Error processing {docx_path}: {str(e)}")
        return 0

def main():
    """Main function to process all DOCX files."""
    
    if not os.path.exists(DOCX_DIR):
        print(f"Error: Directory {DOCX_DIR} not found!")
        return
    
    # Get all DOCX files
    docx_files = sorted([f for f in os.listdir(DOCX_DIR) if f.endswith('.docx')])
    
    if not docx_files:
        print(f"No DOCX files found in {DOCX_DIR}")
        return
    
    print(f"Found {len(docx_files)} DOCX files to process\n")
    
    total_replacements = 0
    
    for docx_file in docx_files:
        input_path = os.path.join(DOCX_DIR, docx_file)
        # Create output filename with "_updated" suffix
        output_filename = docx_file.replace('.docx', '_updated.docx')
        output_path = os.path.join(DOCX_DIR, output_filename)
        
        print(f"\nProcessing: {docx_file}")
        replacements = replace_top_institutions_with_where_to_study(input_path, output_path)
        total_replacements += replacements
    
    print(f"\n{'='*60}")
    print(f"Total replacements made: {total_replacements}")
    print(f"Updated files saved with '_updated' suffix in {DOCX_DIR}")
    print(f"{'='*60}")

if __name__ == "__main__":
    main()
