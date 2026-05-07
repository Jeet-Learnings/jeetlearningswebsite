#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = Document('public/extradata/Combined file - Hospitality and Tourism.docx')

# Map table indices to careers
career_tables = {
    'airhostess': 2,
    'culinary_arts': 5,
    'event_planner': 8,
    'hotel_management': 11,
    'travel_and_tourism': 14,
    'wedding_planner': 17
}

hospitality_institutions = {}

for career, table_idx in career_tables.items():
    if table_idx < len(doc.tables):
        table = doc.tables[table_idx]
        
        # Get all text from the table
        all_text = ""
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + "\n"
        
        # Parse the text
        lines = [line.strip() for line in all_text.split('\n') if line.strip()]
        
        institutions = {
            'government': [],
            'private': [],
            'online': []
        }
        
        current_type = None
        for line in lines:
            line_lower = line.lower()
            
            if line_lower == 'government':
                current_type = 'government'
            elif line_lower == 'private':
                current_type = 'private'
            elif line_lower == 'online':
                current_type = 'online'
            elif line_lower in ['career opportunities', 'conventional', 'new-age', 'remote', 'top institutions']:
                current_type = None
            elif current_type and line and line not in ['Government', 'Private', 'Online']:
                institutions[current_type].append(line)
        
        hospitality_institutions[career] = institutions

print(json.dumps(hospitality_institutions, indent=2, ensure_ascii=False))
