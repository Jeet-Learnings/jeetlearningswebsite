from docx import Document

doc = Document('public/extradata/Combined file - Agriculture.docx')

# Agricultural Engineering pathways table (index 3)
table = doc.tables[3]

print("Extracted Agricultural Engineering Pathways:")
print("=" * 100)
print()

# Get the content from the table
for row_idx, row in enumerate(table.rows):
    cell_text = row.cells[0].text.strip()
    if row_idx == 0:
        # Skip header
        continue
    
    # Parse each pathway
    lines = cell_text.split('\n')
    for line in lines:
        line = line.strip()
        if line:
            print(f'"{line}",')

print()
print("=" * 100)
print("For Agri-Business Management:")
print("=" * 100)

# Agri-Business Management pathways table (index 0)
table = doc.tables[0]

for row_idx, row in enumerate(table.rows):
    cell_text = row.cells[0].text.strip()
    if row_idx == 0:
        # Skip header
        continue
    
    # Parse each pathway
    lines = cell_text.split('\n')
    for line in lines:
        line = line.strip()
        if line:
            print(f'"{line}",')
