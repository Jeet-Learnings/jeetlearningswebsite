#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = Document('public/extradata/Combined file - Hospitality and Tourism.docx')

# Extract table 2 (Air Hostess institutions)
print("=== TABLE 2 (Air Hostess) ===")
table = doc.tables[2]
for row in table.rows:
    for cell in row.cells:
        text = cell.text.strip()
        if text:
            print(repr(text[:150]))

print("\n=== TABLE 5 (Culinary Arts) ===")
table = doc.tables[5]
for row in table.rows:
    for cell in row.cells:
        text = cell.text.strip()
        if text:
            print(repr(text[:150]))
