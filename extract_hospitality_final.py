#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = Document('public/extradata/Combined file - Hospitality and Tourism.docx')

# Map table indices to careers
# Based on the output: Tables 2, 5, 8, 11, 14, 16+ contain institutions
career_tables = {
    'airhostess': 2,
    'culinary_arts': 5,
    'event_planner': 8,
    'hotel_management': 11,
    'travel_and_tourism': 14,
    'wedding_planner': 16
}

hospitality_institutions = {}

for career, table_idx in career_tables.items():
    if table_idx < len(doc.tables):
        table = doc.tables[table_idx]
        institutions = {
            'government': [],
            'private': [],
            'online': []
        }
        
        current_type = None
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                
                if text.lower() == 'government':
                    current_type = 'government'
                elif text.lower() == 'private':
                    current_type = 'private'
                elif text.lower() == 'online':
                    current_type = 'online'
                elif text.lower() in ['career opportunities', 'conventional', 'new-age', 'remote']:
                    current_type = None
                elif current_type and text and text not in ['Government', 'Private', 'Online']:
                    institutions[current_type].append(text)
        
        hospitality_institutions[career] = institutions

print(json.dumps(hospitality_institutions, indent=2, ensure_ascii=False))
