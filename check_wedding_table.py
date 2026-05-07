#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = Document('public/extradata/Combined file - Hospitality and Tourism.docx')
print(f"Total tables: {len(doc.tables)}")

# Check table 16 and 17
for idx in [16, 17]:
    if idx < len(doc.tables):
        print(f"\n=== TABLE {idx} ===")
        table = doc.tables[idx]
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    print(f"Row {row_idx}: {text[:100]}")
                    if row_idx > 5:
                        break
