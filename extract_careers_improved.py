import docx
import json
import glob
import os
from typing import Dict, List, Any
from docx.table import Table
from docx.text.paragraph import Paragraph

def get_document_elements(doc):
    """
    Get all elements (paragraphs and tables) in document order.
    """
    elements = []
    
    # Get parent element
    parent = doc.element.body
    
    for child in parent:
        if child.tag.endswith('p'):  # Paragraph
            # Find the paragraph object
            for para in doc.paragraphs:
                if para._element == child:
                    elements.append(('paragraph', para))
                    break
        elif child.tag.endswith('tbl'):  # Table
            # Find the table object
            for table in doc.tables:
                if table._element == child:
                    elements.append(('table', table))
                    break
    
    return elements

def extract_table_data(table):
    """
    Extract data from a table.
    """
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            # Get cell text, preserving line breaks
            cell_text = cell.text.strip()
            row_data.append(cell_text)
        table_data.append(row_data)
    return table_data

def is_career_title(para):
    """
    Check if a paragraph is a career title (bold and short).
    """
    text = para.text.strip()
    if not text or len(text) > 100:
        return False
    
    # Check if any run is bold
    is_bold = any(run.bold for run in para.runs) if para.runs else False
    
    # Additional checks: not starting with bullet points or numbers
    if text.startswith('•') or text.startswith('-') or (text[0].isdigit() and '.' in text[:5]):
        return False
    
    return is_bold

def extract_career_data(doc):
    """
    Extract all career options with their associated tables.
    """
    elements = get_document_elements(doc)
    careers = []
    current_career = None
    
    for elem_type, elem in elements:
        if elem_type == 'paragraph':
            if is_career_title(elem):
                # Save previous career
                if current_career:
                    careers.append(current_career)
                
                # Start new career
                current_career = {
                    'title': elem.text.strip(),
                    'tables': [],
                    'table_count': 0
                }
        
        elif elem_type == 'table' and current_career:
            # Add table to current career
            table_data = extract_table_data(elem)
            current_career['tables'].append(table_data)
            current_career['table_count'] += 1
    
    # Save last career
    if current_career:
        careers.append(current_career)
    
    return careers

def parse_career_tables(career):
    """
    Parse tables to extract structured information about a career.
    Common table patterns in the documents.
    """
    parsed_data = {
        'title': career['title'],
        'sections': {}
    }
    
    for table_idx, table in enumerate(career['tables']):
        # Try to identify table type by first row
        if not table or not table[0]:
            continue
        
        first_cell = table[0][0].lower() if table[0] else ''
        
        # Store raw table data
        section_name = f'table_{table_idx + 1}'
        
        # Try to identify section by keywords
        if 'pathway' in first_cell or 'career path' in first_cell:
            section_name = 'career_pathways'
        elif 'eligibility' in first_cell or 'qualification' in first_cell:
            section_name = 'eligibility'
        elif 'skill' in first_cell:
            section_name = 'skills'
        elif 'salary' in first_cell or 'package' in first_cell:
            section_name = 'salary'
        elif 'institute' in first_cell or 'college' in first_cell or 'university' in first_cell:
            section_name = 'institutes'
        elif 'entrance' in first_cell or 'exam' in first_cell:
            section_name = 'entrance_exams'
        elif 'job' in first_cell or 'role' in first_cell or 'position' in first_cell:
            section_name = 'job_roles'
        elif 'scope' in first_cell or 'future' in first_cell:
            section_name = 'scope'
        elif 'what is' in first_cell or 'overview' in first_cell:
            section_name = 'overview'
        
        parsed_data['sections'][section_name] = table
    
    return parsed_data

def main():
    """
    Extract career data from all DOCX files.
    """
    all_data = {}
    
    docx_files = glob.glob('public/extradata/*.docx')
    print(f"Found {len(docx_files)} DOCX files to process\n")
    
    total_careers = 0
    
    for filepath in sorted(docx_files):
        try:
            print(f"Processing: {os.path.basename(filepath)}")
            doc = docx.Document(filepath)
            filename = os.path.basename(filepath)
            
            # Extract career data
            careers = extract_career_data(doc)
            
            # Parse each career's tables
            parsed_careers = [parse_career_tables(career) for career in careers]
            
            # Clean category name
            category_name = (filename
                .replace('.docx', '')
                .replace('Combined File - ', '')
                .replace('Combined file - ', '')
                .replace('Combine File - ', '')
                .replace('Comined File - ', ''))
            
            all_data[category_name] = {
                'filename': filename,
                'total_careers': len(careers),
                'careers': parsed_careers
            }
            
            total_careers += len(careers)
            
            print(f"  ✓ Extracted {len(careers)} careers")
            for career in careers:
                print(f"    - {career['title']} ({career['table_count']} tables)")
            print()
            
        except Exception as e:
            print(f"  ✗ Error reading {filepath}: {e}\n")
            import traceback
            traceback.print_exc()
    
    # Save to JSON
    output_file = 'extracted_careers_complete.json'
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_data, f, indent=2, ensure_ascii=False)
    
    print(f"\n{'='*70}")
    print(f"✓ Extraction Complete!")
    print(f"{'='*70}")
    print(f"Total categories: {len(all_data)}")
    print(f"Total careers extracted: {total_careers}")
    print(f"Output file: {output_file}")
    print(f"{'='*70}\n")
    
    # Print detailed summary
    print("=== DETAILED SUMMARY ===\n")
    for category, data in sorted(all_data.items()):
        print(f"\n{category.upper()}")
        print(f"{'-' * len(category)}")
        print(f"Total careers: {data['total_careers']}\n")
        for career in data['careers']:
            sections = list(career['sections'].keys())
            print(f"  • {career['title']}")
            print(f"    Sections: {', '.join(sections)}")
        print()

if __name__ == '__main__':
    main()
