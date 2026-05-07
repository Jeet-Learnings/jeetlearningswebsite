import docx
import sys

def read_docx(path):
    doc = docx.Document(path)
    print("PARAGRAPHS:")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"P[{i}] {text}")
            if i > 20: break
            
    print("\nTABLES:")
    for t_idx, table in enumerate(doc.tables):
        print(f"Table {t_idx}")
        for r_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip().replace("\n", " | "))
            print(f"  Row {r_idx}: " + " || ".join(row_data))
        if t_idx > 2:
            break

if __name__ == "__main__":
    read_docx("public/extradata/Combined File - Account and Finance.docx")
