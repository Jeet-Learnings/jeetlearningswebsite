#!/usr/bin/env python3
"""
Convert all DOCX files to TXT format for dynamic content extraction
"""

from docx import Document
import os
from pathlib import Path

def convert_docx_to_txt(docx_path, txt_path):
    """Convert a single DOCX file to TXT"""
    try:
        doc = Document(docx_path)
        
        with open(txt_path, 'w', encoding='utf-8') as txt_file:
            for para in doc.paragraphs:
                if para.text.strip():
                    txt_file.write(para.text + '\n')
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            txt_file.write(cell.text + '\n')
        
        print(f"✓ Converted: {docx_path} → {txt_path}")
        return True
    except Exception as e:
        print(f"✗ Error converting {docx_path}: {e}")
        return False

def main():
    docx_dir = Path('public/extradata')
    txt_dir = Path('public/content-data')
    
    # Create output directory if it doesn't exist
    txt_dir.mkdir(parents=True, exist_ok=True)
    
    # Find all DOCX files
    docx_files = list(docx_dir.glob('*.docx'))
    
    if not docx_files:
        print(f"No DOCX files found in {docx_dir}")
        return
    
    print(f"Found {len(docx_files)} DOCX files")
    print("=" * 80)
    
    success_count = 0
    for docx_file in sorted(docx_files):
        txt_file = txt_dir / (docx_file.stem + '.txt')
        if convert_docx_to_txt(str(docx_file), str(txt_file)):
            success_count += 1
    
    print("=" * 80)
    print(f"Successfully converted {success_count}/{len(docx_files)} files")

if __name__ == '__main__':
    main()
