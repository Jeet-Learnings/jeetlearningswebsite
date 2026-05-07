#!/usr/bin/env python3
"""
Advanced extraction script that properly parses DOCX structure to extract
career names and their corresponding institutions data.
"""

import os
import re
import json
from docx import Document

DOCX_DIR = "public/extradata"
OUTPUT_FILE = "all_careers_institutions_complete.json"

def extract_careers_from_docx(docx_path):
    """
    Extract careers and institutions by reading document structure.
    Career names appear as paragraphs, followed by tables with data.
    """
    try:
        doc = Document(docx_path)
        careers_dict = {}
        
        current_career = None
        para_idx = 0
        
        # First pass: identify career names (they appear as standalone paragraphs)
        career_names = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and not text.startswith('Pathway') and not text.startswith('Market') and not text.startswith('Where'):
                # Check if this looks like a career name (short, no special formatting)
                if len(text) < 100 and '\n' not in text:
                    # Avoid common non-career text
                    if not any(keyword in text.lower() for keyword in ['salary', 'snapshot', 'institutions', 'opportunities', 'challenges', 'skills']):
                        career_names.append(text)
        
        # Now extract institutions for each career from tables
        table_idx = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Check if this is a career name
            if text in career_names:
                current_career = text
                # Look at the next table for this career's data
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    institutions = extract_institutions_from_table(table)
                    if any(institutions.values()):
                        careers_dict[current_career] = institutions
                    table_idx += 1
        
        return careers_dict
        
    except Exception as e:
        print(f"Error processing {docx_path}: {str(e)}")
        return {}

def extract_institutions_from_table(table):
    """Extract institutions from a table."""
    institutions = {
        "government": [],
        "private": [],
        "online": []
    }
    
    # Convert table to text
    table_text = ""
    for row in table.rows:
        for cell in row.cells:
            table_text += cell.text + "\n"
    
    # Look for institutions section
    if 'Where to Study?' not in table_text and 'Top Institutions' not in table_text:
        return institutions
    
    # Parse institutions
    current_category = None
    for line in table_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        if line.lower().startswith('government:'):
            current_category = 'government'
            content = re.sub(r'^government:\s*', '', line, flags=re.IGNORECASE).strip()
            if content:
                institutions['government'].extend([i.strip() for i in content.split(',')])
        
        elif line.lower().startswith('private:'):
            current_category = 'private'
            content = re.sub(r'^private:\s*', '', line, flags=re.IGNORECASE).strip()
            if content:
                institutions['private'].extend([i.strip() for i in content.split(',')])
        
        elif line.lower().startswith('online:'):
            current_category = 'online'
            content = re.sub(r'^online:\s*', '', line, flags=re.IGNORECASE).strip()
            if content:
                institutions['online'].extend([i.strip() for i in content.split(',')])
        
        elif current_category and line and not any(line.lower().startswith(cat + ':') for cat in ['government', 'private', 'online']):
            if not any(keyword in line.lower() for keyword in ['where', 'top', 'institutions', 'study']):
                institutions[current_category].append(line)
    
    return institutions

def generate_typescript_format(career_name, institutions):
    """Generate TypeScript format for institutions."""
    content_lines = []
    
    if institutions['government']:
        gov_str = "Government: " + ", ".join(institutions['government'])
        content_lines.append(f'"{gov_str}"')
    
    if institutions['private']:
        priv_str = "Private: " + ", ".join(institutions['private'])
        content_lines.append(f'"{priv_str}"')
    
    if institutions['online']:
        online_str = "Online: " + ", ".join(institutions['online'])
        content_lines.append(f'"{online_str}"')
    
    if not content_lines:
        content_lines.append('"Top institutions."')
    
    return ",\n    ".join(content_lines)

def main():
    """Main function."""
    
    if not os.path.exists(DOCX_DIR):
        print(f"Error: Directory {DOCX_DIR} not found!")
        return
    
    # Get all original DOCX files
    docx_files = sorted([f for f in os.listdir(DOCX_DIR) 
                        if f.endswith('.docx') and not f.endswith('_updated.docx')])
    
    if not docx_files:
        print(f"No DOCX files found in {DOCX_DIR}")
        return
    
    print(f"Found {len(docx_files)} DOCX files\n")
    
    all_careers = {}
    
    for docx_file in docx_files:
        input_path = os.path.join(DOCX_DIR, docx_file)
        print(f"Processing: {docx_file}")
        
        careers_data = extract_careers_from_docx(input_path)
        
        if careers_data:
            print(f"  ✓ Extracted {len(careers_data)} careers")
            all_careers.update(careers_data)
        else:
            print(f"  (No data found)")
    
    # Save to JSON
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        json.dump(all_careers, f, indent=2, ensure_ascii=False)
    
    print(f"\n{'='*60}")
    print(f"Total careers: {len(all_careers)}")
    print(f"Saved to: {OUTPUT_FILE}")
    print(f"{'='*60}\n")
    
    # Generate TypeScript snippets
    ts_output_file = "institutions_typescript_snippets.ts"
    with open(ts_output_file, 'w', encoding='utf-8') as f:
        f.write("// Auto-generated institutions data for all careers\n")
        f.write("// Copy and paste the relevant sections into your data files\n\n")
        
        for career_name in sorted(all_careers.keys()):
            institutions = all_careers[career_name]
            ts_content = generate_typescript_format(career_name, institutions)
            
            f.write(f"// {career_name}\n")
            f.write(f"// institutions section:\n")
            f.write(f"{{\n")
            f.write(f"  id: \"institutions\",\n")
            f.write(f"  title: \"Where to Study?\",\n")
            f.write(f"  icon: \"Building2\",\n")
            f.write(f"  description: \"Top institutions across India.\",\n")
            f.write(f"  color: BLUE,\n")
            f.write(f"  content: [\n")
            f.write(f"    {ts_content}\n")
            f.write(f"  ]\n")
            f.write(f"}},\n\n")
    
    print(f"TypeScript snippets saved to: {ts_output_file}")
    
    # Print summary
    print(f"\nCareers extracted:")
    for i, (career_name, institutions) in enumerate(sorted(all_careers.items()), 1):
        gov = len(institutions.get('government', []))
        priv = len(institutions.get('private', []))
        online = len(institutions.get('online', []))
        print(f"{i:2d}. {career_name:40s} | Gov: {gov:2d} | Priv: {priv:2d} | Online: {online:2d}")

if __name__ == "__main__":
    main()
