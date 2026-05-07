#!/usr/bin/env python3
"""
Script to extract "Top Institutions" from DOCX files and replace with "Where to Study?" sections.
Handles multiple careers per document (20+) stored in tables.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX_DIR = "public/extradata"

def extract_and_replace_institutions_in_text(text):
    """
    Extract "Top Institutions" section and replace with "Where to Study?" section.
    Returns the modified text.
    """
    # Pattern to match "Top Institutions" section
    # This pattern looks for "Top Institutions" followed by content until the next section
    pattern = r'Top Institutions\s*(.*?)(?=\n(?:Career Opportunities|Challenges|Future|Skills to Build|Where Are the Jobs|$))'
    
    def replace_func(match):
        institutions_content = match.group(1).strip()
        # Replace "Top Institutions" with "Where to Study?"
        return f"Where to Study?\n{institutions_content}"
    
    modified_text = re.sub(pattern, replace_func, text, flags=re.DOTALL | re.IGNORECASE)
    return modified_text

def process_docx_file(docx_path, output_path):
    """
    Process a DOCX file, find and replace "Top Institutions" with "Where to Study?" in tables.
    """
    try:
        doc = Document(docx_path)
        replacements_made = 0
        
        # Process tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Get cell text
                    cell_text = cell.text
                    
                    if 'Top Institutions' in cell_text:
                        # Replace "Top Institutions" with "Where to Study?"
                        modified_text = cell_text.replace('Top Institutions', 'Where to Study?')
                        
                        # Clear cell and add modified text
                        cell.text = modified_text
                        replacements_made += 1
                        print(f"  ✓ Table {table_idx}, Row {row_idx}, Cell {cell_idx}: Replaced 'Top Institutions'")
        
        # Also process paragraphs (in case content is not in tables)
        for para_idx, para in enumerate(doc.paragraphs):
            if 'Top Institutions' in para.text:
                para.text = para.text.replace('Top Institutions', 'Where to Study?')
                replacements_made += 1
                print(f"  ✓ Paragraph {para_idx}: Replaced 'Top Institutions'")
        
        # Save the modified document
        doc.save(output_path)
        return replacements_made
        
    except Exception as e:
        print(f"✗ Error processing {docx_path}: {str(e)}")
        import traceback
        traceback.print_exc()
        return 0

def main():
    """Main function to process all DOCX files."""
    
    if not os.path.exists(DOCX_DIR):
        print(f"Error: Directory {DOCX_DIR} not found!")
        return
    
    # Get all DOCX files
    docx_files = sorted([f for f in os.listdir(DOCX_DIR) if f.endswith('.docx')])
    
    if not docx_files:
        print(f"No DOCX files found in {DOCX_DIR}")
        return
    
    print(f"Found {len(docx_files)} DOCX files to process\n")
    
    total_replacements = 0
    
    for docx_file in docx_files:
        input_path = os.path.join(DOCX_DIR, docx_file)
        # Create output filename with "_updated" suffix
        output_filename = docx_file.replace('.docx', '_updated.docx')
        output_path = os.path.join(DOCX_DIR, output_filename)
        
        print(f"\nProcessing: {docx_file}")
        replacements = process_docx_file(input_path, output_path)
        total_replacements += replacements
        
        if replacements > 0:
            print(f"✓ Saved to: {output_filename}")
        else:
            print(f"  (No 'Top Institutions' found to replace)")
    
    print(f"\n{'='*60}")
    print(f"Total replacements made: {total_replacements}")
    print(f"Updated files saved with '_updated' suffix in {DOCX_DIR}")
    print(f"{'='*60}")

if __name__ == "__main__":
    main()
