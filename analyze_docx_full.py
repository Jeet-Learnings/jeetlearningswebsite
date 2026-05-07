#!/usr/bin/env python3
"""
Script to analyze the full content of DOCX files.
"""

import os
from pathlib import Path
from docx import Document

DOCX_DIR = "public/extradata"

def analyze_docx_full(docx_path):
    """Analyze a DOCX file and print all content."""
    try:
        doc = Document(docx_path)
        
        print(f"\n{'='*80}")
        print(f"File: {os.path.basename(docx_path)}")
        print(f"{'='*80}")
        
        # Print all paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[{i}] {para.text}")
        
        print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
        
        # Also check tables
        if doc.tables:
            print(f"\nTables found: {len(doc.tables)}")
            for t_idx, table in enumerate(doc.tables):
                print(f"\nTable {t_idx}:")
                for row in table.rows[:5]:  # First 5 rows
                    print([cell.text for cell in row.cells])
        
    except Exception as e:
        print(f"Error analyzing {docx_path}: {str(e)}")

def main():
    """Main function."""
    
    if not os.path.exists(DOCX_DIR):
        print(f"Error: Directory {DOCX_DIR} not found!")
        return
    
    # Get first DOCX file to analyze
    docx_files = sorted([f for f in os.listdir(DOCX_DIR) if f.endswith('.docx')])
    
    if not docx_files:
        print(f"No DOCX files found in {DOCX_DIR}")
        return
    
    # Analyze first file
    first_file = os.path.join(DOCX_DIR, docx_files[0])
    analyze_docx_full(first_file)

if __name__ == "__main__":
    main()
