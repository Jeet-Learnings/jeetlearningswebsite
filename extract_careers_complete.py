import docx
import json
import glob
import os
from typing import Dict, List, Any

def extract_career_data(doc) -> List[Dict[str, Any]]:
    """
    Extract all career options from a DOCX document.
    Career titles are identified by bold formatting.
    """
    careers = []
    current_career = None
    current_section = None
    current_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            continue
        
        # Check if this is a career title (bold paragraph)
        is_bold = any(run.bold for run in para.runs) if para.runs else False
        
        # If it's a bold paragraph and not too long, it's likely a career title
        if is_bold and len(text) < 100 and not text.startswith('•'):
            # Save previous career if exists
            if current_career:
                if current_section and current_content:
                    current_career['sections'][current_section] = '\n'.join(current_content)
                careers.append(current_career)
            
            # Start new career
            current_career = {
                'title': text,
                'sections': {},
                'tables': []
            }
            current_section = None
            current_content = []
        
        # Check if this is a section header (common patterns)
        elif current_career and is_section_header(text):
            # Save previous section content
            if current_section and current_content:
                current_career['sections'][current_section] = '\n'.join(current_content)
            
            current_section = text
            current_content = []
        
        # Regular content
        elif current_career:
            if current_section:
                current_content.append(text)
            else:
                # Content before any section header
                if 'introduction' not in current_career['sections']:
                    current_career['sections']['introduction'] = text
                else:
                    current_career['sections']['introduction'] += '\n' + text
    
    # Save last career
    if current_career:
        if current_section and current_content:
            current_career['sections'][current_section] = '\n'.join(current_content)
        careers.append(current_career)
    
    # Extract tables and associate them with careers
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            table_data.append(row_data)
        
        # Try to associate table with the most recent career
        if careers:
            careers[-1]['tables'].append({
                'index': table_idx,
                'data': table_data
            })
    
    return careers

def is_section_header(text: str) -> bool:
    """
    Identify if a text is likely a section header.
    Common section headers in career documents.
    """
    section_keywords = [
        'what is', 'what are', 'overview', 'introduction',
        'eligibility', 'qualification', 'requirements',
        'skills', 'competencies', 'abilities',
        'career path', 'career options', 'opportunities',
        'salary', 'compensation', 'earnings', 'package',
        'scope', 'future', 'prospects',
        'courses', 'programs', 'degrees',
        'institutes', 'colleges', 'universities',
        'entrance exam', 'admission', 'selection',
        'job roles', 'positions', 'designations',
        'industries', 'sectors', 'domains',
        'challenges', 'difficulties',
        'advantages', 'benefits', 'pros',
        'disadvantages', 'cons',
        'how to become', 'pathway', 'roadmap',
        'top recruiters', 'employers', 'companies',
        'work environment', 'day in life',
        'certifications', 'licenses',
        'specializations', 'branches',
        'duration', 'time period',
        'fees', 'cost', 'investment'
    ]
    
    text_lower = text.lower()
    
    # Check if text matches common section patterns
    if any(keyword in text_lower for keyword in section_keywords):
        return True
    
    # Check if it ends with colon (common for headers)
    if text.endswith(':'):
        return True
    
    # Check if it's a numbered/lettered header
    if text[0].isdigit() and '.' in text[:5]:
        return True
    
    return False

def main():
    """
    Extract career data from all DOCX files in public/extradata/
    """
    all_data = {}
    
    docx_files = glob.glob('public/extradata/*.docx')
    print(f"Found {len(docx_files)} DOCX files to process\n")
    
    for filepath in docx_files:
        try:
            print(f"Processing: {filepath}")
            doc = docx.Document(filepath)
            filename = os.path.basename(filepath)
            
            # Extract career data
            careers = extract_career_data(doc)
            
            # Store in dictionary
            category_name = filename.replace('.docx', '').replace('Combined File - ', '').replace('Combined file - ', '').replace('Combine File - ', '').replace('Comined File - ', '')
            
            all_data[category_name] = {
                'filename': filename,
                'total_careers': len(careers),
                'careers': careers
            }
            
            print(f"  ✓ Extracted {len(careers)} careers")
            for career in careers:
                print(f"    - {career['title']} ({len(career['sections'])} sections, {len(career['tables'])} tables)")
            print()
            
        except Exception as e:
            print(f"  ✗ Error reading {filepath}: {e}\n")
    
    # Save to JSON
    output_file = 'extracted_careers_complete.json'
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_data, f, indent=2, ensure_ascii=False)
    
    print(f"\n{'='*60}")
    print(f"Extraction complete!")
    print(f"Total categories: {len(all_data)}")
    print(f"Total careers: {sum(cat['total_careers'] for cat in all_data.values())}")
    print(f"Output saved to: {output_file}")
    print(f"{'='*60}")
    
    # Print summary
    print("\n=== SUMMARY BY CATEGORY ===\n")
    for category, data in sorted(all_data.items()):
        print(f"{category}: {data['total_careers']} careers")
        for career in data['careers']:
            print(f"  • {career['title']}")
        print()

if __name__ == '__main__':
    main()
