#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import json
import sys

# Set output encoding
sys.stdout.reconfigure(encoding='utf-8')

# Load the hospitality DOCX file
doc = Document('public/extradata/Combined file - Hospitality and Tourism.docx')

# Extract all text from tables
hospitality_data = {}

for table_idx, table in enumerate(doc.tables):
    print(f"\n=== TABLE {table_idx} ===", file=sys.stderr)
    rows = table.rows
    
    # Get all text from this table
    table_text = []
    for row in rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text:
                table_text.append(text)
    
    # Print first 30 lines
    for i, text in enumerate(table_text[:30]):
        try:
            print(f"Line {i}: {text[:100]}", file=sys.stderr)
        except:
            pass

print(f"\nTotal tables: {len(doc.tables)}", file=sys.stderr)
