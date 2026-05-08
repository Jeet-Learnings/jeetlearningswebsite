from docx import Document
import json

doc = Document('public/extradata/Combined - Science, Maths and Engineering_updated_updated.docx')

# Extract career titles
careers = []
for para in doc.paragraphs:
    if para.text.strip() and para.text.strip() not in careers:
        careers.append(para.text.strip())

print(f"Found {len(careers)} careers:")
for i, career in enumerate(careers):
    print(f"{i+1}. {career}")

# Extract table content
print(f"\n\nFound {len(doc.tables)} tables")

# Group tables by career (3 tables per career typically)
tables_per_career = len(doc.tables) // len(careers)
print(f"Approximately {tables_per_career} tables per career")

# Extract first career's data as example
print("\n=== ASTRONOMY DATA ===")
for t_idx in range(min(3, len(doc.tables))):
    table = doc.tables[t_idx]
    print(f"\nTable {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  [{r_idx},{c_idx}]: {text[:300]}")
