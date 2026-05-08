from docx import Document
import json
import re

doc = Document('public/extradata/Combined - Science, Maths and Engineering_updated_updated.docx')

# Extract career titles
careers = []
for para in doc.paragraphs:
    if para.text.strip() and para.text.strip() not in careers:
        careers.append(para.text.strip())

print(f"Found {len(careers)} careers")

# Extract all table data
all_data = {}

for career_idx, career in enumerate(careers):
    career_slug = career.lower().replace(" ", "_").replace("&", "and")
    print(f"\n=== {career} ({career_slug}) ===")
    
    # Each career has 3 tables
    table_start = career_idx * 3
    table_end = table_start + 3
    
    career_data = {
        "name": career,
        "slug": career_slug,
        "pathways": "",
        "market": "",
        "jobs": "",
        "institutions": "",
        "opportunities": ""
    }
    
    for t_idx in range(table_start, min(table_end, len(doc.tables))):
        table = doc.tables[t_idx]
        
        # Extract table content
        table_text = ""
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_text += text + "\n"
        
        # Categorize by table content
        if "Career Pathways" in table_text or "Pathway A" in table_text:
            career_data["pathways"] = table_text
            print(f"Table {t_idx}: Career Pathways")
        elif "Market Snapshot" in table_text or "Salary Snapshot" in table_text:
            career_data["market"] = table_text
            print(f"Table {t_idx}: Market Snapshot")
        elif "Where to Study" in table_text or "Career Opportunities" in table_text:
            career_data["institutions"] = table_text
            print(f"Table {t_idx}: Institutions & Opportunities")
        else:
            print(f"Table {t_idx}: Other content")
    
    all_data[career_slug] = career_data

# Save to JSON for inspection
with open('science_eng_data.json', 'w', encoding='utf-8') as f:
    json.dump(all_data, f, indent=2, ensure_ascii=False)

print("\n\nData saved to science_eng_data.json")

# Print sample
print("\n=== SAMPLE: Astronomy ===")
print(json.dumps(all_data.get("astronomy", {}), indent=2, ensure_ascii=False)[:1000])
