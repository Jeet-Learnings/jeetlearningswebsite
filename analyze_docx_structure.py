#!/usr/bin/env python3
"""
Script to analyze the structure of DOCX files to understand how "Top Institutions" is formatted.
"""

import os
from pathlib import Path
from docx import Document

DOCX_DIR = "public/extradata"

def analyze_docx(docx_path):
    """Analyze a DOCX file and print its structure."""
    try:
        doc = Document(docx_path)
        
        print(f"\n{'='*80}")
        print(f"File: {os.path.basename(docx_path)}")
        print(f"{'='*80}")
        
        # Print first 100 paragraphs to understand structure
        for i, para in enumerate(doc.paragraphs[:100]):
            if para.text.strip():
                # Print paragraph with style info
                style = para.style.name if para.style else "No Style"
                print(f"[{i}] ({style}) {para.text[:100]}")
                
                # Look for "Top Institutions" or similar
                if 'institution' in para.text.lower() or 'where' in para.text.lower():
                    print(f"     ^^^ FOUND RELEVANT TEXT ^^^")
        
        print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
        
    except Exception as e:
        print(f"Error analyzing {docx_path}: {str(e)}")

def main():
    """Main function."""
    
    if not os.path.exists(DOCX_DIR):
        print(f"Error: Directory {DOCX_DIR} not found!")
        return
    
    # Get first DOCX file to analyze
    docx_files = sorted([f for f in os.listdir(DOCX_DIR) if f.endswith('.docx')])
    
    if not docx_files:
        print(f"No DOCX files found in {DOCX_DIR}")
        return
    
    # Analyze first file
    first_file = os.path.join(DOCX_DIR, docx_files[0])
    analyze_docx(first_file)

if __name__ == "__main__":
    main()
