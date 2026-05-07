#!/usr/bin/env python3
"""
Script to directly update TypeScript career data files with institutions data.
Reads from DOCX files and updates the corresponding career entries in TS files.
"""

import os
import re
import json
from docx import Document

DOCX_DIR = "public/extradata"

# Map of DOCX files to their corresponding TypeScript data files
DOCX_TO_TS_MAP = {
    "Combined file - Bio Science and Research.docx": "app/data/careerPageData.ts",
    "Combined File - Business Management.docx": "app/data/careerPageData.ts",
    "Combined File - Education and Training.docx": "app/data/careerPageData.ts",
    "Combined File - Environment.docx": "app/data/careerPageData.ts",
    "Combined File - Health Science.docx": "app/data/careerPageData.ts",
    "Combined File - Human and Social Sciences.docx": "app/data/careerPageData.ts",
    "Combined File - Government Services.docx": "app/data/careerPageData.ts",
    "Combined File - Legal Services.docx": "app/data/careerPageData.ts",
    "Combined File - Logistics and Transportation.docx": "app/data/careerPageData.ts",
    "Combined File - Manufacturing.docx": "app/data/careerPageData.ts",
    "Combined File - Marketing and Advertising.docx": "app/data/careerPageData.ts",
    "Combined File - Media and Communication.docx": "app/data/careerPageData.ts",
    "Combined File - Sports and Physical Activities.docx": "app/data/careerPageData.ts",
    "Combine File - Information Technology.docx": "app/data/careerPageData.ts",
    "Combined - Science, Maths and Engineering.docx": "app/data/careerPageData.ts",
    "Comined File - Public Safety and Security.docx": "app/data/careerPageData.ts",
}

def extract_all_institutions_from_docx(docx_path):
    """Extract all institutions data from a DOCX file."""
    try:
        doc = Document(docx_path)
        all_data = {}
        
        # Get all text to understand structure
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        # Extract from tables
        for table_idx, table in enumerate(doc.tables):
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + "\n"
            
            # Look for institutions section
            if 'Where to Study?' in table_text or 'Top Institutions' in table_text:
                # Parse institutions
                institutions = {
                    "government": [],
                    "private": [],
                    "online": []
                }
                
                current_category = None
                for line in table_text.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    
                    if line.lower().startswith('government:'):
                        current_category = 'government'
                        content = re.sub(r'^government:\s*', '', line, flags=re.IGNORECASE).strip()
                        if content:
                            institutions['government'].extend([i.strip() for i in content.split(',')])
                    
                    elif line.lower().startswith('private:'):
                        current_category = 'private'
                        content = re.sub(r'^private:\s*', '', line, flags=re.IGNORECASE).strip()
                        if content:
                            institutions['private'].extend([i.strip() for i in content.split(',')])
                    
                    elif line.lower().startswith('online:'):
                        current_category = 'online'
                        content = re.sub(r'^online:\s*', '', line, flags=re.IGNORECASE).strip()
                        if content:
                            institutions['online'].extend([i.strip() for i in content.split(',')])
                    
                    elif current_category and line and not any(line.lower().startswith(cat + ':') for cat in ['government', 'private', 'online']):
                        if not any(keyword in line.lower() for keyword in ['where', 'top', 'institutions', 'study', 'career', 'opportunities']):
                            institutions[current_category].append(line)
                
                if any(institutions.values()):
                    all_data[f"table_{table_idx}"] = institutions
        
        return all_data
        
    except Exception as e:
        print(f"Error: {str(e)}")
        return {}

def generate_institutions_content(institutions):
    """Generate TypeScript content for institutions."""
    lines = []
    
    if institutions.get('government'):
        gov_str = "Government: " + ", ".join(institutions['government'][:4])  # Limit to 4
        lines.append(f'"{gov_str}"')
    
    if institutions.get('private'):
        priv_str = "Private: " + ", ".join(institutions['private'][:4])  # Limit to 4
        lines.append(f'"{priv_str}"')
    
    if institutions.get('online'):
        online_str = "Online: " + ", ".join(institutions['online'][:4])  # Limit to 4
        lines.append(f'"{online_str}"')
    
    if not lines:
        lines.append('"Top institutions."')
    
    return ",\n        ".join(lines)

def main():
    """Main function."""
    
    print("Extracting institutions data from DOCX files...\n")
    
    all_institutions = {}
    
    for docx_file in os.listdir(DOCX_DIR):
        if not docx_file.endswith('.docx') or docx_file.endswith('_updated.docx'):
            continue
        
        docx_path = os.path.join(DOCX_DIR, docx_file)
        print(f"Processing: {docx_file}")
        
        institutions_data = extract_all_institutions_from_docx(docx_path)
        
        if institutions_data:
            print(f"  ✓ Found {len(institutions_data)} institution sections")
            all_institutions[docx_file] = institutions_data
            
            # Print sample
            for key, data in list(institutions_data.items())[:1]:
                print(f"    Sample: {key}")
                if data.get('government'):
                    print(f"      Government: {data['government'][:2]}")
                if data.get('private'):
                    print(f"      Private: {data['private'][:2]}")
                if data.get('online'):
                    print(f"      Online: {data['online'][:2]}")
        else:
            print(f"  (No institutions found)")
    
    # Save extracted data
    with open("docx_institutions_extracted.json", 'w', encoding='utf-8') as f:
        json.dump(all_institutions, f, indent=2, ensure_ascii=False)
    
    print(f"\n{'='*60}")
    print(f"Extracted data saved to: docx_institutions_extracted.json")
    print(f"Total DOCX files processed: {len(all_institutions)}")
    print(f"{'='*60}")
    
    # Generate TypeScript snippets
    print(f"\nGenerating TypeScript snippets...\n")
    
    ts_output = []
    ts_output.append("// Auto-generated institutions data from DOCX files")
    ts_output.append("// Use these snippets to update career data files\n")
    
    for docx_file, institutions_dict in all_institutions.items():
        ts_output.append(f"// From: {docx_file}")
        for key, institutions in institutions_dict.items():
            content = generate_institutions_content(institutions)
            ts_output.append("{")
            ts_output.append('  id: "institutions",')
            ts_output.append('  title: "Where to Study?",')
            ts_output.append('  icon: "Building2",')
            ts_output.append('  description: "Top institutions across India.",')
            ts_output.append('  color: BLUE,')
            ts_output.append('  content: [')
            ts_output.append(f'    {content}')
            ts_output.append('  ]')
            ts_output.append("},\n")
    
    with open("docx_institutions_typescript.ts", 'w', encoding='utf-8') as f:
        f.write("\n".join(ts_output))
    
    print(f"✓ TypeScript snippets saved to: docx_institutions_typescript.ts")

if __name__ == "__main__":
    main()
